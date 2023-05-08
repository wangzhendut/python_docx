from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START,WD_ORIENTATION
from docx.shared import Inches
from docx.shared import Pt  #磅
from docx.shared import RGBColor

#创建默认的空白docx文档
document=Document()
#打开已有文档
# document = Document('existing-document-file.docx')

#---------------------------添加内容-------------------------
#添加标题
document.add_heading('一级标题')
document.add_heading('二级标题', level=2)

#在文档末尾添加段落
paragraph = document.add_paragraph('这是通过add_paragraph添加的段落')#style='ListBullet'
#在某段前添加段落
prior_paragraph = paragraph.insert_paragraph_before('在段落前插入段落')
# 加粗.bold和斜体.italic
run = paragraph.add_run('    加粗asfasdadsadasdadadadadasdasdadasdasdasd')
run.bold = True

# or
# paragraph.add_run('加粗').bold = True
# E.g
# paragraph = document.add_paragraph()
# paragraph.add_run('Lorem ipsum ')#style="Emphasis"
# paragraph.add_run('dolor').bold = True
# paragraph.add_run(' sit amet.')

# 设置对齐方式
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#设置缩进
# paragraph_format.left_indent = Inches(2)
#首行缩进
paragraph_format.first_line_indent = Pt(2)
# 段落间距
paragraph_format.space_before = Pt(18)#段前
paragraph_format.space_after  = Pt(18)#段后
# 行距
paragraph_format.line_spacing = Pt(18)
# 设置字体和大小，下划线，字体颜色
font = run.font
font.name = '宋体'
font.size = Pt(12)
font.underline = True
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
# 添加分页符
document.add_page_break()

# 添加表
table = document.add_table(rows=2, cols=2)
# 访问单个单元格
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'
# 访问行
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'
# 对行列循环
# for row in table.rows:
#     for cell in row.cells:
#         print(cell.text)
# 查询行列数目
row_count = len(table.rows)
col_count = len(table.columns)
# 添加行
row = table.add_row()

# 添加图片
document.add_picture('D:\script\python_docx\微信图片_20201011171923.jpg')
# -------------------------设置页面布局----------------------
#访问sections
sections = document.sections
section = sections[0]
current_section = document.sections[-1]
new_section = document.add_section(WD_SECTION_START.ODD_PAGE)
section.start_type = WD_SECTION_START.ODD_PAGE
# 页面尺寸和方向
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
# 页边距
section.left_margin = Inches(1.5)
section.right_margin = Inches(1)

#设置页眉和页脚
section = document.sections[0]
header = section.header
aragraph = header.paragraphs[0]
# paragraph.text = "Title of my document"

#保存文档
document.save('test.docx')
